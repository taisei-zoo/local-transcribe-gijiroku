# document_preprocessor.py
# -*- coding: utf-8 -*-
"""V3用の文書前処理。

対応:
- txt / md: そのまま読み込み
- docx: 段落と表を抽出
- xlsx: シートごとにMarkdown表として抽出
- pdf: テキストPDFのみ抽出。スキャンPDFは警告を返す。

V3ではGUIで警告を見られるため、V4のような
「文字抽出できなかったファイル.json」は標準出力しない。
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None  # type: ignore

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover
    load_workbook = None  # type: ignore

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None  # type: ignore


TEXT_EXTENSIONS = {".txt", ".md"}
DOCUMENT_EXTENSIONS = {".docx", ".xlsx", ".pdf"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCUMENT_EXTENSIONS


@dataclass
class PreprocessResult:
    source_path: Path
    kind: str
    text: str
    warnings: list[str]


DEFAULT_PREPROCESS_CONFIG = {
    "word_extract_tables_as_markdown": True,
    "excel_max_sheets_per_book": 30,
    "excel_max_rows_per_sheet": 5000,
    "excel_max_cols_per_sheet": 80,
    "excel_max_cells_per_sheet": 80000,
    "excel_max_total_cells_per_book": 200000,
    "excel_max_cell_chars": 2000,
    "excel_include_hidden_sheets": False,
    "pdf_min_extracted_chars": 100,
}


def merged_config(config: dict | None) -> dict:
    merged = DEFAULT_PREPROCESS_CONFIG.copy()
    if isinstance(config, dict):
        merged.update(config)
    return merged


def safe_filename(stem: str) -> str:
    for ch in ['\\', '/', ':', '*', '?', '"', '<', '>', '|']:
        stem = stem.replace(ch, '_')
    return stem.strip() or "file"


def unique_path(path: Path) -> Path:
    if not path.exists():
        return path
    i = 2
    while True:
        candidate = path.with_name(f"{path.stem}_{i}{path.suffix}")
        if not candidate.exists():
            return candidate
        i += 1


def natural_key(path: Path):
    return [int(t) if t.isdigit() else t.lower() for t in re.split(r"(\d+)", path.name)]


def read_text_lossy(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8-sig", errors="replace")
    except Exception:
        return path.read_text(encoding="cp932", errors="replace")


def _escape_md(value: Any) -> str:
    text = "" if value is None else str(value)
    text = text.replace("\r", " ").replace("\n", " ").replace("|", "｜")
    return text.strip()


def markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    max_cols = max(len(r) for r in rows)
    norm = [(r + [""] * (max_cols - len(r))) for r in rows]
    header = norm[0]
    lines = ["| " + " | ".join(_escape_md(c) for c in header) + " |"]
    lines.append("| " + " | ".join(["---"] * max_cols) + " |")
    for row in norm[1:]:
        lines.append("| " + " | ".join(_escape_md(c) for c in row) + " |")
    return "\n".join(lines)


def cell_to_text(value: Any, max_chars: int) -> str:
    if value is None:
        return ""
    text = str(value).strip().replace("\r", " ").replace("\n", " ")
    if len(text) > max_chars:
        return text[:max_chars] + "…（セル文字数上限により省略）"
    return text


def trim_empty_edges(rows: list[list[str]]) -> list[list[str]]:
    while rows and not any(c.strip() for c in rows[-1]):
        rows.pop()
    if not rows:
        return []
    max_cols = max(len(r) for r in rows)
    last_nonempty = -1
    for c in range(max_cols):
        if any(c < len(r) and r[c].strip() for r in rows):
            last_nonempty = c
    if last_nonempty < 0:
        return []
    return [r[: last_nonempty + 1] for r in rows]


def extract_docx(path: Path, config: dict | None = None) -> tuple[str, list[str]]:
    cfg = merged_config(config)
    warnings: list[str] = []
    if Document is None:
        raise RuntimeError("python-docx がインストールされていません。")

    doc = Document(str(path))
    lines: list[str] = [f"# Word抽出: {path.name}", ""]

    para_count = 0
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            para_count += 1
            lines.append(text)
            lines.append("")

    if bool(cfg.get("word_extract_tables_as_markdown", True)):
        for t_idx, table in enumerate(doc.tables, start=1):
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                lines.append(f"## 表{t_idx}")
                lines.append(markdown_table(rows))
                lines.append("")
    elif doc.tables:
        warnings.append(f"{path.name}: Word表のMarkdown抽出が無効のため、表は抽出していません。")

    if para_count == 0 and not doc.tables:
        warnings.append(f"{path.name}: 本文・表を抽出できませんでした。")

    return "\n".join(lines).strip(), warnings


def extract_xlsx(path: Path, config: dict | None = None) -> tuple[str, list[str]]:
    cfg = merged_config(config)
    warnings: list[str] = []
    if load_workbook is None:
        raise RuntimeError("openpyxl がインストールされていません。")

    max_sheets = int(cfg.get("excel_max_sheets_per_book", 30))
    max_rows = int(cfg.get("excel_max_rows_per_sheet", 5000))
    max_cols = int(cfg.get("excel_max_cols_per_sheet", 80))
    max_cells_sheet = int(cfg.get("excel_max_cells_per_sheet", 80000))
    max_cells_book = int(cfg.get("excel_max_total_cells_per_book", 200000))
    max_cell_chars = int(cfg.get("excel_max_cell_chars", 2000))
    include_hidden = bool(cfg.get("excel_include_hidden_sheets", False))

    wb = load_workbook(str(path), read_only=True, data_only=True)
    lines: list[str] = [f"# Excel抽出: {path.name}", ""]
    total_cells = 0
    processed_sheets = 0

    try:
        for ws in wb.worksheets:
            if processed_sheets >= max_sheets:
                warnings.append(f"{path.name}: シート数が上限{max_sheets}を超えたため、以降のシートは省略しました。")
                break
            if not include_hidden and getattr(ws, "sheet_state", "visible") != "visible":
                warnings.append(f"{path.name}: 非表示シート「{ws.title}」は省略しました。")
                continue

            sheet_rows_limit = min(max_rows, ws.max_row or 0)
            sheet_cols_limit = min(max_cols, ws.max_column or 0)
            if sheet_rows_limit <= 0 or sheet_cols_limit <= 0:
                continue

            if sheet_rows_limit * sheet_cols_limit > max_cells_sheet:
                allowed_rows = max(1, max_cells_sheet // max(1, sheet_cols_limit))
                warnings.append(f"{path.name}: シート「{ws.title}」は1シートセル上限{max_cells_sheet}により{allowed_rows}行まで抽出しました。")
                sheet_rows_limit = min(sheet_rows_limit, allowed_rows)

            if total_cells + sheet_rows_limit * sheet_cols_limit > max_cells_book:
                remaining = max_cells_book - total_cells
                if remaining <= 0:
                    warnings.append(f"{path.name}: ブック全体セル上限{max_cells_book}に到達したため、以降のシートは省略しました。")
                    break
                allowed_rows = max(1, remaining // max(1, sheet_cols_limit))
                warnings.append(f"{path.name}: ブック全体セル上限{max_cells_book}に近いため、シート「{ws.title}」は{allowed_rows}行まで抽出しました。")
                sheet_rows_limit = min(sheet_rows_limit, allowed_rows)

            rows: list[list[str]] = []
            for row in ws.iter_rows(
                min_row=1,
                max_row=sheet_rows_limit,
                min_col=1,
                max_col=sheet_cols_limit,
                values_only=True,
            ):
                rows.append([cell_to_text(v, max_cell_chars) for v in row])
            rows = trim_empty_edges(rows)
            if rows:
                processed_sheets += 1
                total_cells += sheet_rows_limit * sheet_cols_limit
                lines.append(f"## シート: {ws.title}")
                lines.append(markdown_table(rows))
                lines.append("")

            if (ws.max_row or 0) > max_rows or (ws.max_column or 0) > max_cols:
                warnings.append(f"{path.name}: シート「{ws.title}」は行列上限（{max_rows}行・{max_cols}列）により一部のみ抽出しました。")
    finally:
        try:
            wb.close()
        except Exception:
            pass

    if processed_sheets == 0:
        warnings.append(f"{path.name}: 抽出できるシートがありませんでした。")

    return "\n".join(lines).strip(), warnings


def extract_pdf(path: Path, config: dict | None = None) -> tuple[str, list[str]]:
    cfg = merged_config(config)
    warnings: list[str] = []
    if PdfReader is None:
        raise RuntimeError("pypdf がインストールされていません。")

    min_chars = int(cfg.get("pdf_min_extracted_chars", 100))
    reader = PdfReader(str(path))
    pages: list[str] = []
    for idx, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            pages.append(f"## ページ{idx}\n{text.strip()}")

    joined = "\n\n".join(pages).strip()
    if len(joined) < min_chars:
        warnings.append(f"{path.name}: PDFから十分なテキストを抽出できませんでした。スキャンPDFまたは画像PDFの可能性があります。")
        return "", warnings

    return f"# PDF抽出: {path.name}\n\n{joined}", warnings


def preprocess_file(path: Path, config: dict | None = None) -> PreprocessResult:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return PreprocessResult(path, "text", read_text_lossy(path), [])
    if suffix == ".docx":
        text, warnings = extract_docx(path, config)
        return PreprocessResult(path, "docx", text, warnings)
    if suffix == ".xlsx":
        text, warnings = extract_xlsx(path, config)
        return PreprocessResult(path, "xlsx", text, warnings)
    if suffix == ".pdf":
        text, warnings = extract_pdf(path, config)
        return PreprocessResult(path, "pdf", text, warnings)
    return PreprocessResult(path, "unsupported", "", [f"{path.name}: 対応していないファイル形式です。"])


def write_extracted_text(result: PreprocessResult, out_dir: Path) -> Path | None:
    if not result.text.strip():
        return None
    out_dir.mkdir(parents=True, exist_ok=True)
    suffix = "txt" if result.kind == "text" else result.kind
    out_path = unique_path(out_dir / f"{safe_filename(result.source_path.stem)}_{suffix}_抽出.txt")
    out_path.write_text(result.text, encoding="utf-8", newline="\n")
    return out_path
